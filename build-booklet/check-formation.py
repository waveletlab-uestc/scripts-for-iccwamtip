#!/usr/bin/env python3

import os
import glob
import re

from docx import Document


def check_title(paragraph):
    pass


def check_authors(paragraph):
    authors = paragraph.text
    if re.search(r'\s{2,}', authors):
        print("duplicated spaces")
    if re.search(r'，', authors):
        print("full width version ',' ('，')")


def check_address(paragraphs):
    pass

def end_of_address(paragraph):
    text = paragraph.text.lower()
    return text == '' or 'mail' in text


def check_email(paragraph):
    pass


def check_multi_empty_paragraphs(paragraphs):
    empty_cnt = 0
    for no, p in enumerate(paragraphs, 1):
        if p.text.strip() == '':
            empty_cnt += 1
        else:
            if empty_cnt >= 3:
                #pre = 100*(no/len(paragraphs))
                print(f"{no}/{len(paragraphs)}: continuous empty paragraphs ({empty_cnt})")
                print(p.text[:50])
            empty_cnt = 0


def check_page(doc):
    pass


def check_tables(doc):
    pass


def check_pictures(doc):
    pass


def check_all(fname):
    print(f"checking {fname} ...")
    if not os.path.exists(fname):
        print(f"{fname}: file not exists.")
        return False
    doc = Document(fname)
    paragraphs = doc.paragraphs

    if not paragraphs:
        print(f"{fname}: empty paper")
        return False
    check_title(paragraphs[0])

    if len(paragraphs) < 2:
        print(f"{fname}: paper is not complete, without AUTHORS information")
        return False
    check_authors(doc.paragraphs[1])

    if len(paragraphs) < 3:
        print(f"{fname}: paper is not complete, without ADDRESS information")
        return False
    for no, paragraph in enumerate(paragraphs[2:], 2):
        if end_of_address(paragraph):
            break
    check_address(paragraphs[2:no])

    check_email(paragraphs[no])

    check_multi_empty_paragraphs(paragraphs)
    check_page(doc)
    check_tables(doc)
    check_pictures(doc)


def main():
    for fname in glob.glob('./*.docx'):
        check_all(fname)


if __name__ == '__main__':
    main()
