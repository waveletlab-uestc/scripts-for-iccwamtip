#!/usr/bin/env python3

"""
从论文中提取作者信息
"""

import re
from glob import glob

from docx import Document
from openpyxl import Workbook


def read_authors(fname):
    doc = Document(fname)
    authors = doc.paragraphs[1].text
    def normalize(name):
        return re.sub(r'\s+', ' ', name.strip().upper())
    name_sep = r'[\d\*]\s*|,|，'
    return [normalize(name) for name in re.split(name_sep, authors) if name.strip() != '']


def read_information(fname):
    doc = Document(fname)
    title = doc.paragraphs[0].text.upper()
    first_author = read_authors(fname)[0]
    institution = doc.paragraphs[2].text
    institution = re.sub(r'^\d{1,1}\s*(?=\w)|^\d{1,1}\s*(,\s*\d{1,1}\s*)+', '', doc.paragraphs[2].text)
    for i, p in enumerate(doc.paragraphs):
        if 'abstract:' == p.text.lower().strip():
            abstract = doc.paragraphs[i+1].text
        elif 'keywords:' == p.text.lower().strip():
            keywords = doc.paragraphs[i+1].text
            break
    return [title, first_author, institution, keywords, abstract]


def write_xlsx(data, fname="meta-information.xlsx"):
    workbook = Workbook()
    worksheet = workbook.active
    for row, rcd in enumerate(data, 1):
        for col, cell in enumerate(rcd, 1):
            worksheet.cell(row, col, cell)
    workbook.save(fname)


def main():
    write_xlsx(read_information(fname) for fname in glob('*.docx'))


if __name__ == '__main__':
    main()
