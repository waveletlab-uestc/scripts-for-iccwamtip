#!/usr/bin/env python3

import os
import re
import zipfile
from collections import OrderedDict

from docx import Document

__doc__ = """
从论文中生成论文索引
"""

title = "Content of Proceeding of 17th ICCWAMTIP"
# Python 3.7 之后 dict 的 keys 保证和添加时顺序一致
tracks = OrderedDict([
    ('Theory and Experiment', 'theory-and-experiment.txt'),
    ('Multimedia Technology', 'multimedia-technology.txt'),
    ('Embedded System and Others', 'embedded-system-and-others.txt'),
])

def get_pages(doc):
    """获取 docx 文档的页数"""
    with zipfile.ZipFile(doc) as zf:
        appxml = zf.read('docProps/app.xml').decode()
        return int(re.search('(?<=<Pages>)\s*\d+\s*(?=</Pages>)', appxml).group(0))


def read_file_list(flist):
    with open(flist) as fp:
        return [id_.strip()+'.docx' for id_ in fp.readlines()]


def copy_paragraph(doc, paragraph, author=False):
    p = doc.add_paragraph()
    for run in paragraph.runs:
        r = p.add_run(run.text)
        r.font.superscript = run.font.superscript
        r.font.all_caps = (author == True)


def end_of_content(paragraph):
    return paragraph == '' or 'mail' in paragraph.lower()


def append_content(dst, track_no, no, fname, start):
    src = Document(fname)
    paper_title = src.paragraphs[0].text.strip()
    title = "#{:02}_{:02}: {}{}".format(track_no, no, paper_title, start).upper()
    dst.add_heading(title, 3)
    for no, paragraph in enumerate(src.paragraphs[1:]):
        if end_of_content(paragraph.text):
            break
        copy_paragraph(dst, paragraph, (0 == no))
    dst.add_paragraph()


def add_track(doc, track_no, track_name, flist, start):
    track_title = "Track {:02}: {}".format(track_no, track_name)
    doc.add_heading(track_title, 2)
    for no, fname in enumerate(flist, 1):
        append_content(doc, track_no, no, fname, start)
        start += get_pages(fname)
    return start


def delete_paragraph(paragraph):
    """
    删除段落 paragraph

    WARN: 只是临时的解决方案
    ref: https://github.com/python-openxml/python-docx/issues/33
    """
    p = paragraph._element
    p.getparent().remove(p)
    p._p = p._element = None


def default_template():
    path = os.path
    prefix = path.split(path.abspath(__file__))[0]
    template = path.join(prefix, 'templates', 'paper-index.docx')
    return template if path.exists(template) else None


def format_out(fname='content.docx', template=default_template()):
    doc = Document(template)
    if doc.paragraphs and doc.paragraphs[-1].text == '':
        delete_paragraph(doc.paragraphs[-1])
    doc.add_heading(title, 1)
    start = 1
    for no, name in enumerate(tracks.keys(), 1):
        start = add_track(doc, no, name, read_file_list(tracks[name]), start)
    doc.save(fname)


def main():
    format_out()


if __name__ == '__main__':
    main()
