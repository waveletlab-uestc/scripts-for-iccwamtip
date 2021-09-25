#!/usr/bin/env python3

import re
import zipfile
import os
from collections import OrderedDict
from operator import add
from functools import reduce

from docx import Document
from docx.shared import Inches, Cm
from docx.enum.text import WD_TAB_ALIGNMENT, WD_TAB_LEADER

__doc__ = """
从论文中生成论文作者索引
"""


tracks = OrderedDict({
    'Theory and Experiment': 'theory-and-experiment.txt',
    'Multimedia Technology': 'multimedia-technology.txt',
    'Embedded System and Others': 'embedded-system-and-others.txt',
})


class Author:
    def __init__(self, name, track_id, page_no):
        self.name = name
        self.track_id = track_id
        self.page_no = page_no


def read_file_list(fname):
    with open(fname) as fp:
        return [id_.strip()+'.docx' for id_ in fp.readlines()]


def get_pages(doc):
    """获取 docx 文档的页数"""
    with zipfile.ZipFile(doc) as zf:
        appxml = zf.read('docProps/app.xml').decode()
        return int(re.search('(?<=<Pages>)\s*\d+\s*(?=</Pages>)', appxml).group(0))


def read_authors(fname):
    doc = Document(fname)
    authors = doc.paragraphs[1].text
    def normalize(name):
        return re.sub(r'\s+', ' ', name.strip().upper())
    name_sep = r'[\d\*]\s*|,|，'
    return [normalize(name) for name in re.split(name_sep, authors) if name.strip() != '']


def fetch_author_information():
    authors = []
    start = 1
    for track_id, track_file in enumerate(tracks.values(), 1):
        for id_, fname in enumerate(read_file_list(track_file), 1):
            authors.append([Author(name, f"#{track_id:02}_{id_:02}", start) for name in read_authors(fname)])
            start += get_pages(fname)
    return reduce(add, authors, [])


def default_template():
    path = os.path
    prefix = path.split(path.abspath(__file__))[0]
    template = path.join(prefix, 'templates', 'author-index.docx')
    return template if path.exists(template) else None


def delete_paragraph(paragraph):
    """
    删除段落 paragraph

    WARN: 只是临时的解决方案
    ref: https://github.com/python-openxml/python-docx/issues/33
    """
    p = paragraph._element
    p.getparent().remove(p)
    p._p = p._element = None


def format_out(authors, output="author-index.docx", template=default_template()):
    doc = Document(template)
    if doc.paragraphs and doc.paragraphs[-1].text == '':
        delete_paragraph(doc.paragraphs[-1])
    NAME_LEN = 19
    index = None
    for author in authors:
        if author.name[0] != index:
            index = author.name[0]
            doc.add_heading(index, 2)
        if len(author.name) >= NAME_LEN:
            doc.add_paragraph(author.name)
            p = doc.add_paragraph(f"\t{author.track_id}{' '*5}{author.page_no}")
        else:
            p = doc.add_paragraph(f"{author.name}\t{author.track_id}{' '*5}{author.page_no}")
        p.paragraph_format.tab_stops.add_tab_stop(Cm(4.5), WD_TAB_ALIGNMENT.LEFT, WD_TAB_LEADER.SPACES)

    doc.save(output)


def main():
    authors = fetch_author_information()
    authors.sort(key=lambda author: (author.name, author.track_id))
    format_out(authors)


if __name__ == '__main__':
    main()
